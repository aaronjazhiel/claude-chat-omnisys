from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import uuid
from datetime import datetime

DOCS_DIR = os.path.join(os.path.dirname(__file__), "descargas")
os.makedirs(DOCS_DIR, exist_ok=True)


def generar(contenido, pregunta="Análisis de código"):
    doc = Document()

    # Estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("⚡ OMNISYS — Análisis de Código SICA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    # Metadata
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Consulta: {pregunta}")
    doc.add_paragraph("─" * 60)

    # Contenido — parsear markdown básico
    for linea in contenido.split('\n'):
        linea = linea.rstrip()

        if not linea:
            doc.add_paragraph("")
            continue

        # Headers markdown
        if linea.startswith('### '):
            p = doc.add_heading(linea[4:], level=3)
        elif linea.startswith('## '):
            p = doc.add_heading(linea[3:], level=2)
        elif linea.startswith('# '):
            p = doc.add_heading(linea[2:], level=1)
        elif linea.startswith('- ') or linea.startswith('* '):
            doc.add_paragraph(linea[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', linea):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', linea), style='List Number')
        elif linea.startswith('📄') or linea.startswith('🔧') or linea.startswith('🔍'):
            p = doc.add_paragraph()
            run = p.add_run(linea)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        elif linea.startswith('```') or linea.startswith('   '):
            p = doc.add_paragraph()
            run = p.add_run(linea.replace('```', ''))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            # Negritas **texto**
            partes = re.split(r'(\*\*.*?\*\*)', linea)
            for parte in partes:
                if parte.startswith('**') and parte.endswith('**'):
                    run = p.add_run(parte[2:-2])
                    run.bold = True
                else:
                    p.add_run(parte)

    # Footer
    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    run = footer.add_run("Generado por Omnisys AI — Powered by Claude (Anthropic)")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Guardar
    filename = f"analisis_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(DOCS_DIR, filename)
    doc.save(filepath)
    return filename
